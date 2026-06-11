"""
生成合并优化版简历 - AI Agent 全栈开发工程师
融合三份简历的优点，补充缺失的工作经历、教育背景、GitHub链接等
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.35

# ── 辅助函数 ──
def add_section_title(doc, text):
    """添加章节标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    # 底部边框线
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="8" w:space="1" w:color="333333"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def add_normal(doc, text, bold=False, size=10.5, indent=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, level=0):
    """添加项目符号段落"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("• " + text)
    run.font.size = Pt(10)
    return p

def set_cell_text(cell, text, bold=False, size=9.5, align='left'):
    """设置单元格文字"""
    # 清除默认空段落
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_skill_table(doc, rows_data, col_widths=None):
    """添加技能表格"""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 设置列宽
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    for i, (category, details) in enumerate(rows_data):
        # 类别列
        cat_cell = table.rows[i].cells[0]
        set_cell_text(cat_cell, category, bold=True, size=9.5, align='center')
        set_cell_shading(cat_cell, '2B579A')
        # 类别文字白色
        cat_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cat_cell.width = Cm(2.5)

        # 内容列
        det_cell = table.rows[i].cells[1]
        set_cell_text(det_cell, details, bold=False, size=9.5)
        det_cell.width = Cm(13)

    # 表格边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

# ══════════════════════════════════════════════════════════
# 一、个人信息
# ══════════════════════════════════════════════════════════
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_p.paragraph_format.space_after = Pt(2)
run = name_p.add_run("王 博 宣")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(6)
run = title_p.add_run("AI Agent 全栈开发工程师")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.paragraph_format.space_after = Pt(2)
run = contact_p.add_run("17634013652  |  1507521693@qq.com  |  西安  |  本科 · 应用统计学 ")
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# 一句话 Slogan
slogan_p = doc.add_paragraph()
slogan_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
slogan_p.paragraph_format.space_after = Pt(4)
run = slogan_p.add_run("以自然语言驱动 AI Agent 协作，完成从创意到交付的全栈产品闭环")
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.italic = True

# ══════════════════════════════════════════════════════════
# 二、专业技能
# ══════════════════════════════════════════════════════════
add_section_title(doc, "专业技能")

skill_rows = [
    ("AI & 大模型",
     "DeepSeek / GPT / Claude 多模型调度 · SSE 流式架构 · Prompt 工程（结构化版本管理 + A/B 测试）· "
     "RAG 检索增强生成（混合检索 + Rerank）· Function Calling / Tool Use · Agent 协作链编排 · "
     "MCP 协议工程落地 · LangChain / LangGraph · Transformer / 注意力机制原理"),
    ("前端工程",
     "React 18/19 · TypeScript 严格模式 · Vite 6/7 · Tailwind CSS · Ant Design 6 · Zustand 5 · "
     "ReactFlow 可视化画布 · CodeMirror 6 编辑器 · ECharts 图表"),
    ("后端工程",
     "Python FastAPI / Flask · Node.js Express · Spring Boot（了解）· Supabase (Auth + RLS + Storage) · "
     "RESTful API 设计 · WebSocket 实时通信 · Rate Limit 多级限流 · 微服务架构"),
    ("数据 & 中间件",
     "PostgreSQL (Supabase) · MySQL · Redis（缓存/分布式锁/队列）· Milvus 向量数据库 · "
     "IndexedDB 浏览器端存储 · Kafka / RabbitMQ（了解）· Embedding 向量化"),
    ("工程化 & DevOps",
     "Docker / docker-compose 容器化 · CI/CD (GitHub Actions) · Git / Conventional Commits · "
     "Vitest 单元测试 · PWA 离线方案 · 共享代码架构（Monorepo）· 文档驱动开发"),
    ("核心素养",
     "从 0 到 1 独立交付 · 渐进式架构演进 · 务实技术选型 · AI 辅助高效开发 · "
     "良好的产品思维与用户体验意识 · 持续关注 AI Agent 前沿技术"),
]
add_skill_table(doc, skill_rows)

# ══════════════════════════════════════════════════════════
# 三、工作经历
# ══════════════════════════════════════════════════════════
add_section_title(doc, "工作经历")

# 公司名 + 时间
co_p = doc.add_paragraph()
co_p.paragraph_format.space_after = Pt(2)
run = co_p.add_run("西安寰亿星橙科技有限公司")
run.bold = True
run.font.size = Pt(11)
run = co_p.add_run("  |  AI Agent 全栈开发工程师  |  2025.06 – 至今")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# 技术栈标签
tag_p = doc.add_paragraph()
tag_p.paragraph_format.space_after = Pt(4)
run = tag_p.add_run("Python · FastAPI · LangChain · LlamaIndex · OpenAI / Claude API · React · MySQL · Redis · Docker · MCP")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

work_items = [
    "独立设计并开发多款基于 LLM 的智能 Agent Web 应用，涵盖智能客服、知识问答、自动化工作流等场景，累计服务用户 2000+",
    "设计并实现 Agent 核心架构：多轮对话管理、动态 Prompt 组装、Tool Use 编排、短期/长期记忆管理等模块，支撑复杂交互场景",
    "基于 FastAPI 搭建 Agent 后端服务，通过异步化 + SSE 流式响应 + 连接池扩容 + Redis 请求队列 + Embedding 缓存等手段，将单服务 QPS 从 50 优化至 300+",
    "集成 MCP 协议，开发数据库查询 / 文件系统等自定义 Tool Server，实现 Agent 与外部数据源的动态连接与即时查询",
    "设计 RAG 知识库方案：文档分块策略（512 token + 128 重叠）+ Embedding 向量化 + Milvus 混合检索（向量 + BM25）+ Cohere Rerank，问答准确率提升至 92%+",
    "开发 Prompt 版本管理与 A/B 测试平台，通过系统化 Prompt 调优使特定场景任务完成率提升 35%（p < 0.01）",
]
for item in work_items:
    add_bullet(doc, item)

# ══════════════════════════════════════════════════════════
# 四、项目经验
# ══════════════════════════════════════════════════════════
add_section_title(doc, "项目经验")

# ── 项目一 ──
proj1_title = doc.add_paragraph()
proj1_title.paragraph_format.space_before = Pt(6)
proj1_title.paragraph_format.space_after = Pt(2)
run = proj1_title.add_run("项目一：AI 赋能引擎 —— 自然语言驱动的游戏生成与社区平台")
run.bold = True
run.font.size = Pt(11)

proj1_desc = doc.add_paragraph()
proj1_desc.paragraph_format.space_after = Pt(2)
run = proj1_desc.add_run(
    '用户输入「做个贪吃蛇」 → AI 自动生成可玩的 HTML5 游戏，配合素材管线与社区系统，完成从 Prompt 到发布的完整链路。'
)
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

proj1_tech = doc.add_paragraph()
proj1_tech.paragraph_format.space_after = Pt(2)
run = proj1_tech.add_run(
    "React 18 · TypeScript · Vite 6 · Tailwind CSS · Express · Supabase · DeepSeek API · IndexedDB · ReactFlow"
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

proj1_items = [
    "代码规模：~22,000 行 TypeScript/TSX/CSS/SQL · 138 源文件 · 5 个 Phase 子项目 · 16+ API 端点 · 7 级限流器",
    "工程质量：TypeScript 严格模式零编译错误 · Vitest 24 测试全过 · 端到端 8 项验证覆盖 AI 生成/API/Auth/Supabase/编译",
    "AI 游戏生成：自然语言 → DeepSeek API 流式返回完整 HTML5 单文件游戏（Canvas 渲染引擎），支持 12 套模板 + 自由生成双模式",
    "流式多轮对话：SSE 实时输出 + 版本快照/回退 + AI 自动错误检测与精准修复（零侧效应）",
    "本地素材库：IndexedDB 5 Store · 完整 CRUD · 占位符绑定 · 游戏内素材引用自动解析",
    "可视化工作流：13 种节点类型 · ReactFlow 画布 · 游戏开发管线编排",
    "社区系统：Supabase Auth · Feed 游标分页 · Remix 链 · 点赞/评论 · Express 后端 16+ API",
    "安全设计：游戏代码 iframe srcdoc 沙箱隔离 · API Key 仅存后端环境变量 · Supabase RLS 权限控制",
    "Docker 部署：完整 Dockerfile + docker-compose.yml · Node.js 原生健康检查 · 零额外依赖",
]
for item in proj1_items:
    add_bullet(doc, item)

# ── 项目二 ──
proj2_title = doc.add_paragraph()
proj2_title.paragraph_format.space_before = Pt(10)
proj2_title.paragraph_format.space_after = Pt(2)
run = proj2_title.add_run("项目二：AI 写作引擎 —— 网络小说作者的 AI 写作伴侣")
run.bold = True
run.font.size = Pt(11)

proj2_desc = doc.add_paragraph()
proj2_desc.paragraph_format.space_after = Pt(2)
run = proj2_desc.add_run(
    "覆盖大纲 → 创作 → 修改 → 分支结局的完整辅助工作流，支持 Markdown 编辑器 + AI 智能续写 + 去 AI 味引擎 + 角色伴侣，PWA 桌面化方案。"
)
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

proj2_tech = doc.add_paragraph()
proj2_tech.paragraph_format.space_after = Pt(2)
run = proj2_tech.add_run(
    "React 19 · TypeScript · Vite 7 · Tailwind CSS · Ant Design 6 · CodeMirror 6 · Zustand 5 · Express · ECharts · PWA"
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

proj2_items = [
    "代码规模：35 源文件 · 5,400+ 行 TypeScript/TSX · 独立前端 + 后端工程 · PWA 可安装为桌面应用",
    "文档体系：7 大模块系统化设计文档（项目规划 → 架构设计 → Prompt 工程 → 前端 → 后端 → Agent 调度 → 测试验收）",
    "AI 智能续写：SSE 流式输出 + 光标位置感知 + 上下文自动拼接 + 接受/拒绝/重试交互",
    "大纲分析系统：AI 结构评分 + ECharts 雷达图可视化 + 梗概一键展开完整章纲",
    "去 AI 味引擎：AI 腔特征检测 → 自然化重写 → Diff 前后对比视图 · 多文风可配置",
    "角色伴侣：角色卡片 + 性格对话模拟 + 关系值追踪 + 变化曲线可视化",
    "分支结局：多分支管理 + Git 风格分支树可视化 + 逐章续写保持原作一致性",
    "AI 驱动开发工作流：独立设计 9 专业 Agent 角色协作体系，流水线式交付（需求 → 架构 → 前后端并行 → 审查 → 验证）",
    "共享代码架构：code/shared/src/ 单一代码源，@shared 别名引用，一处修改全项目同步",
]
for item in proj2_items:
    add_bullet(doc, item)

# ── Demo 链接 ──
demo_p = doc.add_paragraph()
demo_p.paragraph_format.space_before = Pt(10)
demo_p.paragraph_format.space_after = Pt(2)
run = demo_p.add_run("🔗 在线 Demo：")
run.bold = True
run.font.size = Pt(10)
run = demo_p.add_run(" ai-writing-engine-production.up.railway.app")
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

# ══════════════════════════════════════════════════════════
# 五、教育背景
# ══════════════════════════════════════════════════════════
add_section_title(doc, "教育背景")

edu_p = doc.add_paragraph()
edu_p.paragraph_format.space_after = Pt(2)
run = edu_p.add_run("塔里木大学")
run.bold = True
run.font.size = Pt(11)
run = edu_p.add_run("  |  应用统计学  |  本科  |  2025.06")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

edu_items = [
    "主修课程：概率论与数理统计、多元统计分析、数据挖掘、Python 程序设计、数据库原理、机器学习",
    "应用统计学背景为 AI 工作提供了扎实的数据敏感度：A/B 测试的显著性检验、RAG 检索效果评估、缓存命中率统计等均直接受益",
]
for item in edu_items:
    add_bullet(doc, item)

# ══════════════════════════════════════════════════════════
# 六、自我评价
# ══════════════════════════════════════════════════════════
add_section_title(doc, "自我评价")

eval_items = [
    "热衷于 AI Agent 技术前沿，具备从 0 到 1 的独立交付能力，擅长将前沿方案转化为可落地的工程实现",
    "动手能力强：一个项目从想法到上线可独立走完全程，同时善于借助 AI Agent 协作提升开发效率",
    "良好的工程习惯：文档先行、代码评审、单元测试、Conventional Commits，注重长期可维护性",
    "务实的技术选型观：优先成熟 SaaS，能用库的绝不自己造，API 兼容层方便切换 LLM 供应商",
    "持续在技术社区分享 Agent 相关的技术实践与项目经验，保持技术敏感度",
]
for item in eval_items:
    add_bullet(doc, item)

# ══════════════════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════════════════
desktop = os.path.join(os.environ['USERPROFILE'], 'Desktop')
output_path = os.path.join(desktop, '个人简历-AI全栈开发工程师-优化版.docx')
try:
    doc.save(output_path)
except PermissionError:
    # File locked by WPS — save with a new name
    import time
    output_path = os.path.join(desktop, f'个人简历-AI全栈开发工程师-优化版-{int(time.time())}.docx')
    doc.save(output_path)
print(f'已保存至: {output_path}')
print(f"✅ 简历已保存至: {output_path}")
